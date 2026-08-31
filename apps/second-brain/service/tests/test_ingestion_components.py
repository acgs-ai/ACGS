import hashlib
import io
import time
from ipaddress import ip_address
from pathlib import Path
from uuid import uuid4

import httpx
from docx import Document as DocxDocument
from pytest import MonkeyPatch, raises

from second_brain.chunking import chunk_document
from second_brain.parsers import ParseFailure, parse_document
from second_brain.providers import FakeEmbeddingProvider, FakeGenerationProvider
from second_brain.storage import (
    FilesystemStorage,
    StoredObjectMismatch,
    object_key,
    sanitize_filename,
)
from second_brain.url_ingest import (
    BoundResponse,
    HttpxBoundTransport,
    SafeUrlError,
    SystemResolver,
    fetch_safe_url,
    validate_url,
    validate_url_syntax,
)


def minimal_pdf(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, value in enumerate(objects, 1):
        offsets.append(len(result))
        result.extend(f"{number} 0 obj\n".encode() + value + b"\nendobj\n")
    xref = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode())
    result.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(result)


def test_filesystem_storage_uses_opaque_keys_and_bounded_atomic_io(tmp_path: Path) -> None:
    storage = FilesystemStorage(tmp_path / "objects", max_bytes=32)
    key = object_key(uuid4(), uuid4(), uuid4())
    stored = storage.write(key, b"private bytes")

    assert stored.key == key
    assert stored.size == 13
    assert storage.read(key) == b"private bytes"
    assert sanitize_filename("../../secret\x00 report?.txt") == "secret report_.txt"
    assert not list((tmp_path / "objects").rglob("*.tmp"))
    with raises(FileExistsError):
        storage.write(key, b"replacement")
    assert storage.read(key) == b"private bytes"

    with raises(ValueError, match="invalid object key"):
        storage.write("../../outside", b"no")
    with raises(ValueError, match="byte limit"):
        storage.write(key, b"x" * 33)
    storage.delete(key)
    assert not (tmp_path / "objects" / key).exists()


def test_filesystem_storage_rejects_symlink_escape_and_tampering(tmp_path: Path) -> None:
    root = tmp_path / "objects"
    outside = tmp_path / "outside"
    outside.mkdir()
    key = object_key(uuid4(), uuid4(), uuid4())
    owner = key.split("/")[0]
    root.mkdir()
    (root / owner).symlink_to(outside, target_is_directory=True)
    storage = FilesystemStorage(root, max_bytes=64)
    with raises(OSError):
        storage.write_partial(key, b"private", hashlib.sha256(b"private").hexdigest())
    assert not list(outside.rglob("*"))

    safe_key = object_key(uuid4(), uuid4(), uuid4())
    expected = hashlib.sha256(b"private").hexdigest()
    storage.write(safe_key, b"private")
    final = root.joinpath(*safe_key.split("/"))
    final.write_bytes(b"tampered")
    with raises(StoredObjectMismatch):
        storage.read(safe_key, expected)


def test_pdf_docx_and_deterministic_chunk_locations() -> None:
    pdf = parse_document(minimal_pdf("PDF provenance evidence"), "pdf", 10_000)
    assert pdf.text == "PDF provenance evidence"
    assert pdf.passages[0].location == {"page_number": 1}

    buffer = io.BytesIO()
    document = DocxDocument()
    document.add_paragraph("First evidence paragraph")
    document.add_paragraph("Second evidence paragraph")
    document.save(buffer)
    docx = parse_document(buffer.getvalue(), "docx", 10_000)
    assert [passage.location for passage in docx.passages] == [
        {"paragraph_number": 1},
        {"paragraph_number": 2},
    ]

    version_id = uuid4()
    first = chunk_document(version_id, docx, max_chunks=10, size=20, overlap=4)
    second = chunk_document(version_id, docx, max_chunks=10, size=20, overlap=4)
    assert first == second
    assert [chunk.ordinal for chunk in first] == list(range(len(first)))
    assert first[0].location["paragraph_number"] == 1


def test_parsers_fail_closed_on_invalid_and_oversized_content() -> None:
    txt = parse_document(b"TXT provenance evidence", "txt", 100)
    assert txt.text == "TXT provenance evidence"
    assert txt.passages[0].location == {"char_start": 0, "char_end": 23}

    with raises(ParseFailure, match="PDF parser rejected"):
        parse_document(b"private-invalid-pdf", "pdf", 100)
    with raises(ParseFailure, match="character limit"):
        parse_document(b"too many characters", "txt", 3)


def test_fake_model_providers_are_deterministic_and_offline(
    monkeypatch: MonkeyPatch,
) -> None:
    def reject_network(*args: object, **kwargs: object) -> None:
        raise AssertionError("fake providers attempted network egress")

    monkeypatch.setattr("socket.socket", reject_network)
    embeddings = FakeEmbeddingProvider(dimensions=4)
    first = embeddings.embed(["bounded evidence"])
    assert first == embeddings.embed(["bounded evidence"])
    assert len(first[0]) == 4
    assert len(FakeEmbeddingProvider(dimensions=64).embed(["bounded evidence"])[0]) == 64
    assert FakeGenerationProvider("grounded fixture").generate("ignored") == "grounded fixture"


class StaticResolver:
    def __init__(self, answers: dict[str, tuple[str, ...]]) -> None:
        self.answers = answers

    def resolve(  # type: ignore[no-untyped-def]
        self, hostname: str, port: int, deadline: float
    ):
        del port, deadline
        return tuple(ip_address(value) for value in self.answers[hostname])


class ScriptedTransport:
    def __init__(self, responses: list[BoundResponse]) -> None:
        self.responses = responses

    def get(  # type: ignore[no-untyped-def]
        self, url: str, host: str, address, timeout: float, max_bytes: int
    ):
        return self.responses.pop(0)


def test_url_validation_blocks_private_metadata_userinfo_and_ports() -> None:
    blocked = StaticResolver(
        {
            "private.test": ("10.0.0.1",),
            "metadata.test": ("169.254.169.254",),
            "ipv6.test": ("fe80::1",),
            "mapped.test": ("::ffff:127.0.0.1",),
            "mixed.test": ("93.184.216.34", "10.0.0.1"),
        }
    )
    for url in (
        "http://private.test",
        "http://metadata.test/latest/meta-data",
        "http://ipv6.test",
        "http://mapped.test",
        "http://mixed.test",
    ):
        with raises(SafeUrlError, match="prohibited"):
            validate_url(url, blocked)
    public = StaticResolver({"public.test": ("93.184.216.34",)})
    with raises(SafeUrlError, match="user information"):
        validate_url("http://user:password@public.test", public)
    with raises(SafeUrlError, match="user information"):
        validate_url("http://@public.test", public)
    with raises(SafeUrlError, match="port"):
        validate_url("https://public.test:8443", public)


def test_url_syntax_validation_does_not_resolve_and_dns_child_obeys_deadline(
    monkeypatch: MonkeyPatch,
) -> None:
    def forbidden_resolution(*args: object, **kwargs: object) -> None:
        del args, kwargs
        raise AssertionError("request-time URL validation attempted DNS")

    monkeypatch.setattr("socket.getaddrinfo", forbidden_resolution)
    assert validate_url_syntax("https://public.test/path?private=query")[0] == (
        "https://public.test/path?private=query"
    )

    resolver = SystemResolver("import time; time.sleep(60)")
    started = time.monotonic()
    with raises(SafeUrlError, match="deadline"):
        resolver.resolve("public.test", 443, started + 0.02)
    assert time.monotonic() - started < 0.5


def test_url_fetch_revalidates_redirects_peer_and_size_without_public_network() -> None:
    public_address = ip_address("93.184.216.34")
    other_public_address = ip_address("93.184.216.35")
    resolver = StaticResolver(
        {
            "one.test": (str(public_address),),
            "two.test": (str(other_public_address),),
            "private.test": ("127.0.0.1",),
        }
    )
    redirect_private = ScriptedTransport(
        [BoundResponse(302, {"location": "http://private.test/secret"}, b"", public_address)]
    )
    with raises(SafeUrlError, match="prohibited"):
        fetch_safe_url("http://one.test", resolver=resolver, transport=redirect_private)

    wrong_peer = ScriptedTransport(
        [BoundResponse(200, {"content-type": "text/plain"}, b"ok", other_public_address)]
    )
    with raises(SafeUrlError, match="peer address"):
        fetch_safe_url("http://one.test", resolver=resolver, transport=wrong_peer)

    missing_peer = ScriptedTransport(
        [BoundResponse(200, {"content-type": "text/plain"}, b"ok", None)]
    )
    with raises(SafeUrlError, match="unavailable"):
        fetch_safe_url("http://one.test", resolver=resolver, transport=missing_peer)

    encoded = ScriptedTransport(
        [
            BoundResponse(
                200,
                {"content-type": "text/plain", "content-encoding": "gzip"},
                b"encoded",
                public_address,
            )
        ]
    )
    with raises(SafeUrlError, match="encoding"):
        fetch_safe_url("http://one.test", resolver=resolver, transport=encoded)

    redirect_userinfo = ScriptedTransport(
        [BoundResponse(302, {"location": "http://@two.test/final"}, b"", public_address)]
    )
    with raises(SafeUrlError, match="user information"):
        fetch_safe_url("http://one.test", resolver=resolver, transport=redirect_userinfo)

    class SlowTransport:
        def get(self, *args: object, **kwargs: object) -> BoundResponse:
            del args, kwargs
            time.sleep(0.02)
            return BoundResponse(200, {"content-type": "text/plain"}, b"ok", public_address)

    with raises(SafeUrlError, match="deadline"):
        fetch_safe_url(
            "http://one.test", resolver=resolver, transport=SlowTransport(), timeout=0.001
        )

    oversized = ScriptedTransport(
        [BoundResponse(200, {"content-type": "text/plain"}, b"12345", public_address)]
    )
    with raises(SafeUrlError, match="byte limit"):
        fetch_safe_url("http://one.test", resolver=resolver, transport=oversized, max_bytes=4)

    redirected = ScriptedTransport(
        [
            BoundResponse(
                302,
                {"location": "http://two.test/final?private=redirect"},
                b"",
                public_address,
            ),
            BoundResponse(200, {"content-type": "text/plain"}, b"evidence", other_public_address),
        ]
    )
    result = fetch_safe_url(
        "http://one.test/start?private=submitted", resolver=resolver, transport=redirected
    )
    assert tuple(result) == (
        b"evidence",
        "text/plain",
        "http://two.test/final?private=redirect",
    )
    assert [hop.uri for hop in result.redirects] == [
        "http://one.test/start?private=submitted",
        "http://two.test/final?private=redirect",
    ]


def test_httpx_transport_pins_numeric_peer_and_preserves_host_and_sni(
    monkeypatch: MonkeyPatch,
) -> None:
    captured: dict[str, object] = {}
    address = ip_address("93.184.216.34")

    class NetworkStream:
        def get_extra_info(self, name: str) -> tuple[str, int] | None:
            return (str(address), 443) if name == "server_addr" else None

    class Response:
        def __init__(self) -> None:
            self.status_code = 200
            self.headers = {"content-type": "text/plain", "content-length": "2"}
            self.extensions = {"network_stream": NetworkStream()}

        def iter_raw(self):  # type: ignore[no-untyped-def]
            yield b"ok"

        def close(self) -> None:
            return None

    class Client:
        def __init__(self, **kwargs: object) -> None:
            captured["client"] = kwargs

        def __enter__(self):  # type: ignore[no-untyped-def]
            return self

        def __exit__(self, *args: object) -> None:
            del args

        def send(self, request, stream: bool):  # type: ignore[no-untyped-def]
            captured["request"] = request
            captured["stream"] = stream
            return Response()

    monkeypatch.setattr("second_brain.url_ingest.httpx.Client", Client)
    response = HttpxBoundTransport().get(
        "https://public.test/evidence", "public.test", address, 1, 100
    )
    request = captured["request"]
    assert isinstance(request, httpx.Request)
    assert response.content == b"ok"
    assert captured["client"] == {"follow_redirects": False, "timeout": 1, "trust_env": False}
    assert request.url.host == str(address)
    assert request.headers["host"] == "public.test"
    assert request.headers["accept-encoding"] == "identity"
    assert request.extensions["sni_hostname"] == b"public.test"
