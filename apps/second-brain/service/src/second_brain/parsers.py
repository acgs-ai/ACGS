import io
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from second_brain.upload_validation import DOCX_MIME, normalize_text, validate_upload


class ParseFailure(ValueError):
    pass


@dataclass(frozen=True)
class Passage:
    text: str
    location: dict[str, Any]


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    passages: tuple[Passage, ...]


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._ignored_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._ignored_depth:
            self._ignored_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth:
            self.parts.append(data)


def _normalize(value: str) -> str:
    return re.sub(r"[ \t]+", " ", value.replace("\r\n", "\n").replace("\r", "\n")).strip()


def _validate_docx_archive(data: bytes, max_chars: int) -> None:
    with ZipFile(io.BytesIO(data)) as archive:
        entries = archive.infolist()
        names = {entry.filename for entry in entries}
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ParseFailure("DOCX parser rejected the source")
        if len(entries) > 10_000 or any(entry.flag_bits & 0x1 for entry in entries):
            raise ParseFailure("DOCX parser rejected the source")
        if sum(entry.file_size for entry in entries) > max(1_000_000, max_chars * 20):
            raise ParseFailure("DOCX archive exceeds the configured expansion limit")


def parse_document(data: bytes, source_type: str, max_chars: int) -> ParsedDocument:
    passages: tuple[Passage, ...]
    try:
        if source_type == "pdf":
            validate_upload("source.pdf", "application/pdf", data, len(data), max_chars)
        elif source_type == "docx":
            validate_upload("source.docx", DOCX_MIME, data, len(data), max_chars)
        if source_type in {"note", "markdown", "txt"}:
            text = normalize_text(data)
            passages = (Passage(text, {"char_start": 0, "char_end": len(text)}),)
        elif source_type == "html":
            extractor = _TextExtractor()
            extractor.feed(data.decode("utf-8"))
            text = _normalize(" ".join(extractor.parts))
            passages = (Passage(text, {"char_start": 0, "char_end": len(text)}),)
        elif source_type == "pdf":
            if not data.startswith(b"%PDF-"):
                raise ParseFailure("PDF parser rejected the source")
            pages: list[Passage] = []
            for page_number, page in enumerate(PdfReader(io.BytesIO(data)).pages, 1):
                page_text = _normalize(page.extract_text() or "")
                if page_text:
                    pages.append(Passage(page_text, {"page_number": page_number}))
            if not pages:
                raise ParseFailure("PDF contains no extractable text")
            passages = tuple(pages)
            text = "\n\n".join(page.text for page in pages)
        elif source_type == "docx":
            _validate_docx_archive(data, max_chars)
            paragraphs = tuple(
                Passage(value, {"paragraph_number": number})
                for number, paragraph in enumerate(DocxDocument(io.BytesIO(data)).paragraphs, 1)
                if (value := _normalize(paragraph.text))
            )
            if not paragraphs:
                raise ParseFailure("DOCX contains no extractable text")
            passages = paragraphs
            text = "\n\n".join(paragraph.text for paragraph in paragraphs)
        else:
            raise ParseFailure("unsupported parser")
    except (
        BadZipFile,
        PackageNotFoundError,
        UnicodeDecodeError,
        OSError,
        PdfReadError,
        ValueError,
    ) as exc:
        if isinstance(exc, ParseFailure):
            raise
        raise ParseFailure(f"{source_type.upper()} parser rejected the source") from exc
    if len(text) > max_chars:
        raise ParseFailure("extracted text exceeds configured character limit")
    if not text:
        raise ParseFailure("source contains no extractable text")
    return ParsedDocument(text, passages)
